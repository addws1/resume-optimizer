"""生成码上飞产品体验报告 DOCX"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.size = Pt(10.5)
style.font.name = 'Microsoft YaHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_run(para, text, bold=False, size=10.5, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return run


def set_spacing(para, before=0, after=0, line_spacing=1.25):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_border_below(para, color='003366'):
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=True, size=20, color=(0, 51, 102))
    set_spacing(p, before=0, after=6)


def add_h1(text):
    p = doc.add_paragraph()
    add_run(p, text, bold=True, size=14, color=(0, 51, 102))
    set_spacing(p, before=16, after=4)
    add_border_below(p)


def add_h2(text):
    p = doc.add_paragraph()
    add_run(p, text, bold=True, size=12, color=(0, 51, 102))
    set_spacing(p, before=12, after=2)


def add_p(text, bold=False, size=10.5):
    p = doc.add_paragraph()
    set_spacing(p, 1, 1, 1.3)
    add_run(p, text, bold=bold, size=size)


def add_empty():
    p = doc.add_paragraph()
    set_spacing(p, 2, 2, 1.0)
    add_run(p, '', size=6)


def add_bold_line(label, value):
    p = doc.add_paragraph()
    set_spacing(p, 2, 2, 1.3)
    add_run(p, label, bold=True, size=10.5)
    add_run(p, value, size=10.5)


def add_bullet(text):
    p = doc.add_paragraph()
    set_spacing(p, 1, 1, 1.25)
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, '· ', size=10.5)
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run(p, part[2:-2], bold=True, size=10.5)
        else:
            add_run(p, part, size=10.5)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, size=10, color=(255, 255, 255))
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '003366')
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            add_run(p, val, size=9.5)


# =============================================

add_title('码上飞产品体验报告')
add_p('体验人：游翔  |  日期：2026-07-22  |  投递岗位：Agent 产品实习生', size=9)

# ---- 1 ----
add_h1('1. 做了什么应用 + 哪里卡住')

add_h2('创建的应用')
add_p('简历收集工具：用户端填写姓名/手机号/邮箱/求职意向、上传简历文件、提交；管理员端查看列表、按求职意向筛选、下载简历文件。')

add_h2('使用场景')
add_p('自己公司/团队招人用，简洁现代风。')

add_h2('生成耗时')
add_p('约 5-10 分钟。')

add_h2('卡住最久的环节')
add_p('打开生成结果的那一刻——页面默认跳转到了管理端的简历列表页，里面有一批看起来像真实数据的假姓名、假简历。表单入口找不到，字段也填不了。')
add_empty()
add_p('最要命的不是技术问题，是我的第一反应——等了这么久，出来一个不能用、看不懂的页面，我立刻开始怀疑自己：是不是需求没说清楚？选的”简洁现代风”影响了功能？是不是不该选”自己公司用”？产品出了 bug，用户却在怀疑自己。这不是交互问题，是心理问题——当一个产品让人从”它可能不行”直接跳到”可能是我自己不行”，用户就走了。')
add_empty()
add_p('后续我跟产品方（AI 客服“小李”）反馈了问题，提出了修改需求——默认首页改为表单页、清理测试假数据。AI 回复说已经改好了，但再次打开时功能并没有完成。更尴尬的是，两次修改之后免费额度就用完了，想继续改也没办法。')
add_empty()
add_p('这一段体验集中暴露了三个问题叠加的致命效果：生成 bug + AI 客服说了不算 + 额度断崖。单独一个还能忍，三个一起出现的时候，用户就不是“想反馈”了，而是“想关掉”。')

# ---- 2 ----
add_h1('2. 谁最需要 + 一句话推荐')

add_h2('目标用户')
add_bullet('产品经理/产品方向学生——脑子里有产品想法、能想清楚功能和流程，但没有开发资源帮自己落地。码上飞让他们直接验证想法，不用等排期。')
add_bullet('学生——课程项目、创业比赛、个人作品集需要一个能跑的 Demo。学校不教编程，找外包太贵，码上飞是零成本的“技术合伙人”。')

add_h2('一句话推荐')
add_p('“以前有个想法，得先找到会写代码的人才敢说；现在你把想法说出来，码上飞直接给你一个能跑的东西。”')

# ---- 3 ----
add_h1('3. 三个改进点（按优先级排序）')

add_h2('P0 — AI 反馈闭环失效：嘴上说改好了，实际没改')
add_bold_line('问题：', '发现 bug 后，我跟产品内置的 AI 客服反馈了修改需求，AI 回复说已经完成修改。但再次打开应用时，功能并没有变化。同时免费额度已经被扣掉了。')
add_bold_line('影响：', '这个问题最严重——不是“第一次做错了”，而是“跟我说改好了结果没改”。用户愿意反馈是给产品第二次机会，但 AI 的虚假反馈把这次信任也消耗掉了。加上额度被扣——用户付出了成本（免费额度也是成本），得到了一个谎言——这个体验比不出 bug 更让人想离开。')
add_bold_line('建议：', 'AI 客服回复“改好了”之前，应该做一个自检——打开生成的链接，验证改动是否真的生效。做不到就不要说“改好了”。如果不能自动验证，至少在回复里加一句“由于是 AI 自动修改，可能会有遗漏，建议您手动检查一下”。诚实比速度重要。')

add_h2('P1 — 生成结果首次体验不可用')
add_bold_line('问题：', '生成完成后，默认打开的链接是管理端列表页，而非应聘者填写的表单页。页面中存在测试假数据（假人名、假简历），导致第一印象是“这东西不能填、不能用”。')
add_bold_line('影响：', '用户核心任务——“拿到一个能用的简历收集工具”——在第一步就失败了。信任崩塔只在一瞬间。')
add_bold_line('建议：', '默认首页应为表单页（应聘者视角），且首次展示前自动清理所有测试数据。')

add_h2('P2 — 生成时间太长')
add_bold_line('问题：', '从提交需求到生成完成需要 5-10 分钟。虽然有进度反馈，但这个等待时间仍然远超用户对“一句话生成应用”的心理预期。')
add_bold_line('影响：', '用户对产品的第一印象是“慢”。生成类工具的核心体验在于“说完就出来”的爽感，十分钟的等待会把这个爽感完全稀释。而且当前面两个问题已经把用户耐心消耗掉之后，长等待就变成了压死骆驼的最后一根稻草。')
add_bold_line('建议：', '两端优化：一是技术侧缩短生成耗时，分步渲染——先出框架让用户看到东西（30 秒内），再逐步加载数据和细节；二是降低预期——生成前提示“预计需要 3-5 分钟，建议先喝杯水”，比让用户干等强。')

# ---- 4 ----
add_h1('4. 同类产品对比')

add_table(
    ['产品', '输出物', '使用体验'],
    [
        ['Claude', '代码/文案/方案', '使用流畅，错误少，响应稳定。但给你的是代码，要自己部署。'],
        ['Cursor', '代码（IDE 内）', '没用过。'],
        ['Coze', 'Agent（对话机器人）', '经常高峰期不可用。输出聊天机器人，不是应用。'],
        ['码上飞', '可运行的应用', '理念最好——直接给能用的应用。但体验还不够可靠：生成慢、无进度、首次遇 bug。'],
    ]
)
add_empty()

add_h2('更喜欢哪个？为什么？')
add_p('诚实地说，日常我还是更偏向 Claude。')
add_empty()
add_p('码上飞想做的事情比 Claude 更进一步——Claude 给你代码，码上飞给你一个直接能跑的完整应用。但当前阶段，Claude 胜在稳定可靠：出错少、响应快、不会让我在第一秒怀疑自己是不是用错了。')
add_empty()
add_p('码上飞让我看到的方向是对的——“说人话 → 得应用”这个体验一旦打磨到足够流畅可靠，它对非技术用户的价值是 Claude 永远做不到的。但现在还差一口气：那口气就是让我第一次打开时，不产生“这什么玩意儿”的念头。')

# ---- 附录 ----
add_h1('附录：生成结果截图')
add_p('（截图 1：生成过程中让你选场景/风格的界面）')
add_p('（截图 2：首次打开时的管理端列表页——含假数据）')
add_p('（截图 3：修改后的表单页——如果产品方改好了的话）')

doc.save('码上飞产品体验报告.docx')
print('done')
