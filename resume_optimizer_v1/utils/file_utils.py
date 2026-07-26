"""
=============================================================================
简历优化助手 · 文件处理工具
=============================================================================
兼容多种编码的文本读取、PDF/DOCX 简历文件解析、文件大小校验。
=============================================================================
"""

import hashlib
import io
from pathlib import Path
from typing import Optional

from config import MAX_FILE_SIZE_MB
from utils.logger import log_error, log_debug


# ── 编码检测优先级（中文优先）──────────────────────────────
_ENCODINGS = ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]


def read_text_file(file_path: str) -> str:
    """
    智能读取文本文件，自动尝试多种编码。

    Args:
        file_path: 文件路径

    Returns:
        解码后的文本内容

    Raises:
        ValueError: 所有编码均失败时抛出
    """
    path = Path(file_path)
    raw_bytes = path.read_bytes()

    for enc in _ENCIDINGS:
        try:
            content = raw_bytes.decode(enc)
            log_debug(f"文件 {path.name} 使用编码 {enc} 解码成功")
            return content
        except (UnicodeDecodeError, LookupError):
            continue

    raise ValueError(f"无法解码文件 {path.name}，已尝试编码：{_ENCIDINGS}")


def read_uploaded_file(uploaded_file) -> str:
    """
    读取 Streamlit UploadedFile 对象，自动检测编码。

    Args:
        uploaded_file: st.file_uploader 返回的文件对象

    Returns:
        解码后的文本内容

    Raises:
        ValueError: 解码失败
    """
    raw_bytes = uploaded_file.read()
    uploaded_file.seek(0)  # 重置指针，允许后续再次读取

    for enc in _ENCIDINGS:
        try:
            return raw_bytes.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue

    raise ValueError(
        f"无法解码文件 {uploaded_file.name}，已尝试编码：{_ENCIDINGS}"
    )


def read_uploaded_bytes(uploaded_file) -> bytes:
    """读取上传文件的原始字节并重置指针"""
    raw = uploaded_file.read()
    uploaded_file.seek(0)
    return raw


def validate_file_size(uploaded_file, max_mb: int = MAX_FILE_SIZE_MB) -> bool:
    """
    校验上传文件大小是否在允许范围内。

    Args:
        uploaded_file: Streamlit UploadedFile
        max_mb: 最大允许的 MB 数

    Returns:
        True 表示文件大小合规
    """
    uploaded_file.seek(0, 2)  # 移动到末尾获取大小
    size_bytes = uploaded_file.tell()
    uploaded_file.seek(0)     # 重置
    max_bytes = max_mb * 1024 * 1024
    if size_bytes > max_bytes:
        return False
    return True


# ══════════════════════════════════════════════════════════════
# PDF 解析（使用 PyPDF2，轻量无 GPU 依赖）
# ══════════════════════════════════════════════════════════════

def parse_pdf(uploaded_file) -> str:
    """
    从上传的 PDF 文件中提取文本。

    依赖 PyPDF2（已在 requirements 中声明）。
    优先使用 pdfplumber（更准确），回退 PyPDF2。

    Args:
        uploaded_file: Streamlit UploadedFile（.pdf）

    Returns:
        提取的纯文本内容
    """
    raw_bytes = read_uploaded_bytes(uploaded_file)

    # 策略 1：pdfplumber（表格、排版保留更好）
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text()
                if txt:
                    text_parts.append(txt)
        if text_parts:
            return "\n\n".join(text_parts)
    except Exception as e:
        log_error("pdf_parse", e, "pdfplumber 解析失败，回退 PyPDF2")

    # 策略 2：PyPDF2（兼容性更广）
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(raw_bytes))
        text_parts = []
        for page in reader.pages:
            txt = page.extract_text()
            if txt:
                text_parts.append(txt)
        return "\n\n".join(text_parts)
    except Exception as e:
        log_error("pdf_parse", e, "PyPDF2 解析也失败")
        raise ValueError(f"PDF 解析失败：{e}")


# ══════════════════════════════════════════════════════════════
# DOCX 解析
# ══════════════════════════════════════════════════════════════

def parse_docx(uploaded_file) -> str:
    """
    从上传的 DOCX 文件中提取文本。

    Args:
        uploaded_file: Streamlit UploadedFile（.docx）

    Returns:
        提取的纯文本内容
    """
    raw_bytes = read_uploaded_bytes(uploaded_file)
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # 同时提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)
    except Exception as e:
        log_error("docx_parse", e, "DOCX 解析失败")
        raise ValueError(f"DOCX 解析失败：{e}")


def parse_resume_file(uploaded_file) -> str:
    """
    自动识别文件类型并提取文本。
    支持：.txt, .md, .pdf, .docx

    Args:
        uploaded_file: Streamlit UploadedFile

    Returns:
        提取的文本内容
    """
    fname = uploaded_file.name.lower()

    if fname.endswith(".pdf"):
        return parse_pdf(uploaded_file)
    elif fname.endswith(".docx"):
        return parse_docx(uploaded_file)
    elif fname.endswith((".txt", ".md")):
        return read_uploaded_file(uploaded_file)
    else:
        # 尝试当文本读取
        return read_uploaded_file(uploaded_file)


# ══════════════════════════════════════════════════════════════
# 内容哈希（文档去重）
# ══════════════════════════════════════════════════════════════

def content_hash(text: str, algo: str = "md5") -> str:
    """
    计算文本内容的哈希值，用于文档去重。

    Args:
        text: 文本内容
        algo: 哈希算法（md5 / sha256）

    Returns:
        十六进制哈希字符串
    """
    h = hashlib.new(algo)
    h.update(text.encode("utf-8", errors="replace"))
    return h.hexdigest()


def compute_file_hash(uploaded_file, algo: str = "md5") -> str:
    """计算上传文件的哈希值"""
    raw = read_uploaded_bytes(uploaded_file)
    h = hashlib.new(algo)
    h.update(raw)
    return h.hexdigest()
