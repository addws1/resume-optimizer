"""
=============================================================================
简历优化助手 · 导出工具
=============================================================================
Markdown / DOCX 文件导出，包含格式化 Word 文档生成。
=============================================================================
"""

import re
from datetime import datetime
from pathlib import Path

from utils.logger import log_error, log_info


# ══════════════════════════════════════════════════════════════
# Markdown 导出
# ══════════════════════════════════════════════════════════════

def export_markdown(content: str, filename: str = "optimized_resume.md") -> str:
    """
    返回 Markdown 内容用于 Streamlit download_button。
    实际写入动作由 Streamlit 的 st.download_button 完成。

    Args:
        content: Markdown 格式文本
        filename: 建议的文件名

    Returns:
        原始内容（供 download_button 的 data 参数使用）
    """
    return content


# ══════════════════════════════════════════════════════════════
# DOCX 导出（复用 md2docx.py 核心逻辑）
# ══════════════════════════════════════════════════════════════

def export_docx_from_markdown(
    md_content: str,
    output_path: str,
    title: str = "简历优化版",
) -> bool:
    """
    将 Markdown 内容转换为格式化 Word 文档。

    复用项目已有的 md2docx.py 排版逻辑。

    Args:
        md_content: Markdown 格式的简历内容
        output_path: 输出 .docx 文件路径
        title: 文档标题

    Returns:
        True 表示导出成功
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── 页面设置 ──
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        # ── 默认字体 ──
        style = doc.styles['Normal']
        style.font.size = Pt(10.5)
        style.font.name = '宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # ── 辅助函数 ──
        def _add_run(para, text, bold=False, size=10.5, color=None, font_name='宋体'):
            run = para.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            return run

        def _set_spacing(para, before=0, after=0, line_spacing=1.35):
            pf = para.paragraph_format
            pf.space_before = Pt(before)
            pf.space_after = Pt(after)
            pf.line_spacing = line_spacing

        def _add_bottom_border(para):
            pPr = para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '003366')
            pBdr.append(bottom)
            pPr.append(pBdr)

        # ── 解析 Markdown 行 ──
        lines = md_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # 空行 / 水平线 / 引用
            if not line.strip() or line.strip() == '---' or line.strip().startswith('>'):
                i += 1
                continue

            # H1
            if line.startswith('# ') and not line.startswith('## '):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_run(p, line[2:].strip(), bold=True, size=22, font_name='微软雅黑')
                _set_spacing(p, before=0, after=4)
                i += 1
                continue

            # H2
            if line.startswith('## '):
                p = doc.add_paragraph()
                _add_run(p, line[3:].strip(), bold=True, size=14, color=(0, 51, 102), font_name='微软雅黑')
                _set_spacing(p, before=14, after=6)
                _add_bottom_border(p)
                i += 1
                continue

            # H3
            if line.startswith('### '):
                p = doc.add_paragraph()
                _add_run(p, line[3:].strip(), bold=True, size=12, color=(0, 51, 102), font_name='微软雅黑')
                _set_spacing(p, before=10, after=2)
                i += 1
                continue

            # 列表项
            if line.strip().startswith('- '):
                text = line.strip()[2:]
                sub = re.match(r'^\*\*(.+?)\*\*[：:]\s*(.*)$', text)
                p = doc.add_paragraph()
                _set_spacing(p, before=1, after=1, line_spacing=1.35)
                p.paragraph_format.left_indent = Cm(0.5)
                _add_run(p, '• ', size=10.5, font_name='宋体')
                if sub:
                    _add_run(p, sub.group(1) + '：', bold=True, size=10.5, font_name='微软雅黑')
                    _add_run(p, sub.group(2), size=10.5, font_name='宋体')
                else:
                    parts = re.split(r'(\*\*.+?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            _add_run(p, part[2:-2], bold=True, size=10.5, font_name='微软雅黑')
                        else:
                            _add_run(p, part, size=10.5, font_name='宋体')
                i += 1
                continue

            # 粗体行
            bm = re.match(r'^\*\*(.+?)\*\*(.*)$', line.strip())
            if bm:
                p = doc.add_paragraph()
                _set_spacing(p, before=1, after=1, line_spacing=1.4)
                label = bm.group(1).strip()
                value = bm.group(2).strip()
                if value.startswith('|'):
                    value = value[1:].strip()
                _add_run(p, label, bold=True, size=10.5, font_name='微软雅黑')
                _add_run(p, value, size=10.5, font_name='宋体')
                i += 1
                continue

            i += 1

        doc.save(output_path)
        log_info(f"DOCX 导出成功: {output_path}")
        return True

    except Exception as e:
        log_error("docx_export", e, f"导出失败: {output_path}")
        return False


def build_export_filename(prefix: str = "export", ext: str = "md") -> str:
    """生成带时间戳的导出文件名"""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{prefix}_{ts}.{ext}"
