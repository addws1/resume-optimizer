"""
将 Markdown 简历转换为格式化的 Word 文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def set_cell_border(paragraph, val='single', sz='4', space='1', color='003366'):
    """给段落设置底边框作为分隔线"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), val)
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_run(paragraph, text, bold=False, size=None, color=None, font_name=None):
    """添加格式化文本块"""
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    p = doc.add_paragraph()
    if level == 1:
        # 主标题 - 姓名
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True, size=22, font_name='微软雅黑')
        set_paragraph_spacing(p, before=0, after=4)
    elif level == 2:
        # 章节标题
        add_run(p, text, bold=True, size=14, color=(0, 51, 102), font_name='微软雅黑')
        set_paragraph_spacing(p, before=14, after=6)
        # 底部边框作为分隔
        set_cell_border(p, val='single', sz='4', space='1', color='003366')
    return p


def add_bold_label_line(doc, label, value):
    """添加「加粗标签：普通内容」格式的段落"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line_spacing=1.4)
    add_run(p, label, bold=True, size=10.5, font_name='微软雅黑')
    add_run(p, value, bold=False, size=10.5, font_name='宋体')
    return p


def add_bullet(doc, text, indent_level=0):
    """添加项目符号段落"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line_spacing=1.35)
    # 手动添加项目符号
    prefix = '• '
    add_run(p, prefix + text, bold=False, size=10.5, font_name='宋体')
    p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
    return p


def add_normal(doc, text, size=10.5, bold=False, indent=0):
    """添加普通段落"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line_spacing=1.35)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    add_run(p, text, bold=bold, size=size, font_name='宋体')
    return p


def build_docx(md_path, docx_path):
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── 默认字体 ──
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 解析 Markdown ──
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过水平线
        if line.strip() == '---':
            i += 1
            continue

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 跳过引用说明
        if line.strip().startswith('>'):
            i += 1
            continue

        # H1 - 姓名
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            add_heading_styled(doc, text, level=1)
            i += 1
            continue

        # H2 - 章节标题
        if line.startswith('## '):
            text = line[3:].strip()
            add_heading_styled(doc, text, level=2)
            i += 1
            continue

        # H3 - 子标题（项目名）
        if line.startswith('### '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=10, after=2)
            add_run(p, text, bold=True, size=12, color=(0, 51, 102), font_name='微软雅黑')
            i += 1
            continue

        # 粗体行（联系人 / 求职意向 / 项目背景）
        if line.strip().startswith('**') and '**' in line.strip()[2:]:
            # 格式：**加粗部分** 后续普通文字
            match = re.match(r'^\*\*(.+?)\*\*(.*)$', line.strip())
            if match:
                label = match.group(1).strip()
                value = match.group(2).strip()
                # 去掉 value 开头的 |
                if value.startswith('|'):
                    value = value[1:].strip()
                add_bold_label_line(doc, label, value)
                i += 1
                continue

        # 列表项（- 开头）
        if line.strip().startswith('- '):
            text = line.strip()[2:]

            # 检查是否为子标题格式（**加粗**：内容）
            sub_match = re.match(r'^\*\*(.+?)\*\*[：:]\s*(.*)$', text)
            if sub_match:
                label = sub_match.group(1).strip()
                content = sub_match.group(2).strip()
                # 作为粗体标签 + 内容
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=1, after=1, line_spacing=1.35)
                p.paragraph_format.left_indent = Cm(0.5)
                add_run(p, '• ', bold=False, size=10.5, font_name='宋体')
                add_run(p, label + '：', bold=True, size=10.5, font_name='微软雅黑')
                add_run(p, content, bold=False, size=10.5, font_name='宋体')
            else:
                # 检查是否需要处理第二层缩进
                indent = 0
                if text.startswith('  '):
                    indent = 0.5
                    text = text.strip()
                # 普通列表项
                # 处理内部加粗
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=1, after=1, line_spacing=1.35)
                p.paragraph_format.left_indent = Cm(0.5)
                add_run(p, '• ', bold=False, size=10.5, font_name='宋体')

                # 解析行内加粗
                parts = re.split(r'(\*\*.+?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        add_run(p, part[2:-2], bold=True, size=10.5, font_name='微软雅黑')
                    else:
                        add_run(p, part, bold=False, size=10.5, font_name='宋体')
            i += 1
            continue

        # 普通文本行（如 URL、描述性文字）
        add_normal(doc, line.strip())
        i += 1
        continue

    # ── 保存 ──
    doc.save(docx_path)
    print(f'Word doc generated: {docx_path}')


if __name__ == '__main__':
    md_path = r'd:\python\游翔_AI产品实习生_简历优化版.md'
    docx_path = r'd:\python\游翔_AI产品实习生_简历优化版.docx'
    build_docx(md_path, docx_path)
