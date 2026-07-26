"""
=============================================================================
简历优化 Agent · DOCX 生成模块
=============================================================================
将优化结果输出为格式化的 Word 文档。
支持基于原始 DOCX 模板生成（保留原字体/样式/边距）。
=============================================================================
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import OUTPUT_DIR
from section_parser import parse_sections_with_titles


# ══════════════════════════════════════════════════════════════
# 字体辅助函数
# ══════════════════════════════════════════════════════════════

def _set_run_font(run, font_name: str = "微软雅黑", size=None,
                  bold: bool = False, color: tuple | None = None,
                  italic: bool = False):
    """
    设置 run 的字体（正确处理中文字体的 w:eastAsia 属性）。
    缺少此设置会导致中文在 Word 中显示为回退字体或方块。
    """
    run.font.name = font_name
    # 设置东亚字体（中文实际生效依赖此属性）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

    if size is not None:
        run.font.size = Pt(size) if isinstance(size, (int, float)) else size
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color) if isinstance(color, tuple) else color
    if italic:
        run.font.italic = True


def _add_styled_paragraph(doc, text: str, font_name: str = "微软雅黑",
                          font_size: int = 11, bold: bool = False,
                          color: tuple | None = None, italic: bool = False,
                          alignment=None, spacing_after: int = 6):
    """添加一个带完整字体设置的段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(spacing_after)
    if text:
        run = p.add_run(text)
        _set_run_font(run, font_name, font_size, bold, color, italic)
    return p


# ══════════════════════════════════════════════════════════════
# 主生成函数
# ══════════════════════════════════════════════════════════════

def generate_docx(
    optimized_content: str,
    self_assessment: str,
    original_text: str = "",
    template_path: Optional[str] = None,
) -> str:
    """
    生成简历优化 DOCX 文件。

    Args:
        optimized_content: 优化后的四栏内容（LLM 输出的 Markdown）
        self_assessment: Agent 自评总结
        original_text: 原始简历文本（用于附录）
        template_path: 原始 DOCX 文件路径（用作模板，保留样式）

    Returns:
        生成的 DOCX 文件路径
    """
    # ── 创建文档（优先用模板）──
    if template_path and os.path.exists(template_path):
        try:
            doc = Document(template_path)
            # 清除模板正文内容，保留样式/页边距/页眉页脚
            body = doc.element.body
            to_remove = []
            for child in body:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('p', 'tbl'):
                    to_remove.append(child)
            for child in to_remove:
                body.remove(child)
        except Exception:
            doc = Document()
            _setup_default_page(doc)
    else:
        doc = Document()
        _setup_default_page(doc)

    # ── 封面标题 ──
    _add_styled_paragraph(
        doc, "简历优化报告",
        font_name="微软雅黑", font_size=22, bold=True,
        color=(0x1a, 0x56, 0xdb),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=4,
    )
    _add_styled_paragraph(
        doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
        font_name="微软雅黑", font_size=10,
        color=(0x66, 0x66, 0x66),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12,
    )
    _add_styled_paragraph(
        doc, "本文档为 AI 优化分析报告，优化后的简历条目请参见「✅ 优化版本」板块。",
        font_name="微软雅黑", font_size=9,
        color=(0x99, 0x99, 0x99),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=6,
    )

    # ── 分割线 ──
    _add_styled_paragraph(doc, "—" * 40, font_size=8, color=(0xcc, 0xcc, 0xcc))

    # ── 优化内容（健壮解析）──
    sections = parse_sections_with_titles(optimized_content)

    for section_title, section_body in sections:
        # 小节标题
        _add_styled_paragraph(
            doc, section_title,
            font_name="黑体", font_size=14, bold=True,
            color=(0x1a, 0x56, 0xdb), spacing_after=4,
        )

        # 小节内容
        for paragraph_text in section_body.strip().split("\n"):
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue
            _add_styled_paragraph(
                doc, paragraph_text,
                font_name="微软雅黑", font_size=11, spacing_after=4,
            )

    # ── 分割线 ──
    _add_styled_paragraph(doc, "—" * 40, font_size=8, color=(0xcc, 0xcc, 0xcc))

    # ── Agent 自评总结 ──
    _add_styled_paragraph(
        doc, "🤖 Agent 自评总结",
        font_name="黑体", font_size=14, bold=True,
        color=(0x1a, 0x56, 0xdb), spacing_after=6,
    )
    if self_assessment:
        _add_styled_paragraph(
            doc, self_assessment,
            font_name="微软雅黑", font_size=11, italic=True, spacing_after=4,
        )

    # ── 附录：原始简历 ──
    if original_text.strip():
        _add_styled_paragraph(doc, "", font_size=6)  # 空行
        _add_styled_paragraph(doc, "—" * 40, font_size=8, color=(0xcc, 0xcc, 0xcc))
        _add_styled_paragraph(
            doc, "📄 附录：原始简历",
            font_name="黑体", font_size=13, bold=True,
            color=(0x66, 0x66, 0x66), spacing_after=6,
        )
        for paragraph_text in original_text.strip().split("\n"):
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue
            _add_styled_paragraph(
                doc, paragraph_text,
                font_name="微软雅黑", font_size=10,
                color=(0x88, 0x88, 0x88), spacing_after=2,
            )

    # ── 保存 ──
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"简历优化_{timestamp}.docx"
    filepath = OUTPUT_DIR / filename
    doc.save(str(filepath))

    return str(filepath)


def generate_clean_resume_docx(resume_md: str) -> str:
    """
    将合成后的干净 Markdown 简历转换为简洁的 DOCX 文件。
    自动去除所有 Markdown 格式符号（**、*、# 等），输出纯文本排版。
    """
    import re

    doc = Document()
    _setup_default_page(doc)

    for line in resume_md.strip().split("\n"):
        line = line.strip()
        if not line:
            continue

        # 一级标题（# 开头，非 ##）→ 姓名
        if line.startswith("# ") and not line.startswith("## "):
            text = _clean_md(line[2:].strip())
            _add_styled_paragraph(
                doc, text,
                font_name="微软雅黑", font_size=18, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=8,
            )
        # 二级标题 → 章节标题
        elif line.startswith("## "):
            _add_styled_paragraph(doc, "", font_size=4)  # 空行
            text = _clean_md(line[3:].strip())
            _add_styled_paragraph(
                doc, text,
                font_name="黑体", font_size=14, bold=True,
                color=(0x1a, 0x56, 0xdb), spacing_after=6,
            )
        # 三级标题 → 项目名称
        elif line.startswith("### "):
            text = _clean_md(line[4:].strip())
            _add_styled_paragraph(
                doc, text,
                font_name="微软雅黑", font_size=12, bold=True, spacing_after=4,
            )
        # 分隔线
        elif line.strip() == "---":
            _add_styled_paragraph(
                doc, "─" * 50, font_size=8, color=(0xcc, 0xcc, 0xcc), spacing_after=2,
            )
        # Bullet 列表
        elif line.startswith("- "):
            text = _clean_md(line[2:].strip())
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(4)
            pf.left_indent = Cm(0.5)
            run = p.add_run(f"• {text}")
            _set_run_font(run, "微软雅黑", 11)
        # 普通段落（去除所有 markdown 符号）
        else:
            text = _clean_md(line)
            _add_styled_paragraph(
                doc, text,
                font_name="微软雅黑", font_size=11, spacing_after=4,
            )

    # 保存
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"简历_可投递_{timestamp}.docx"
    filepath = OUTPUT_DIR / filename
    doc.save(str(filepath))

    return str(filepath)


def _clean_md(text: str) -> str:
    """去除 Markdown 内联格式符号：**bold** → bold，*italic* → italic，`code` → code"""
    import re
    # 去除 **bold**
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 去除 *italic*（单星号包围，不与 ** 冲突）
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
    # 去除 `code`
    text = re.sub(r'`(.+?)`', r'\1', text)
    # 去除残留的孤立 * 号
    text = text.replace('*', '')
    return text


def _setup_default_page(doc: Document):
    """设置默认 A4 页面（当没有模板时使用）"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
