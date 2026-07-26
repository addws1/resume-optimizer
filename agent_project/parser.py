"""
=============================================================================
简历优化 Agent · 文件解析模块（v2.1）
=============================================================================
支持 PDF（三级引擎回退）和 DOCX 的文本提取。

增强：
  - PDF：PyMuPDF → pdfplumber → PyPDF2 三级回退
  - 文本编码：UTF-8 → GBK → GB2312 → GB18030 → Latin-1 五级回退
  - 图片扫描件：明确提示用户粘贴文本
=============================================================================
"""

from pathlib import Path
from typing import Tuple, Optional


class ParseError(Exception):
    """文件解析错误，携带用户可理解的提示"""
    def __init__(self, message: str, user_msg: str):
        super().__init__(message)
        self.user_msg = user_msg  # 给用户看的友好提示


# ══════════════════════════════════════════════════════════════
# 多编码回退
# ══════════════════════════════════════════════════════════════

# 按常见程度排序：UTF-8 → 中文编码 → 兜底
_ENCODING_CHAIN = ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]


def read_text_with_encoding_fallback(data: bytes) -> str:
    """
    用多种编码尝试解码文本，返回第一个成功的结果。

    国内用户上传的文件常为 GBK/GB2312 编码的旧版 Word 导出，
    单靠 UTF-8 会直接报错。此函数提供五级回退。
    """
    for enc in _ENCODING_CHAIN:
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue

    raise ParseError(
        "无法识别文件编码",
        "❌ 文件编码不支持，请将文件另存为 UTF-8 后重试，或直接粘贴简历文本。"
    )


# ══════════════════════════════════════════════════════════════
# 主入口
# ══════════════════════════════════════════════════════════════

def parse_resume_file(file_path: str) -> Tuple[str, str, Optional[str]]:
    """
    解析简历文件，返回 (文本内容, 文件类型标签, 模板路径)。

    Args:
        file_path: PDF 或 DOCX 文件路径

    Returns:
        (extracted_text, label, template_path)
        - label 如 "PDF" 或 "DOCX"
        - template_path: DOCX 返回原始路径（用于模板化生成），PDF 返回 None

    Raises:
        ParseError: 解析失败，携带用户友好提示
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(path)
    elif ext == ".docx":
        return _parse_docx(path)
    else:
        raise ParseError(
            f"不支持的文件格式: {ext}",
            f"不支持的文件格式（{ext}），请上传 PDF 或 DOCX 文件。"
        )


# ══════════════════════════════════════════════════════════════
# PDF 解析（三级回退）
# ══════════════════════════════════════════════════════════════

def _parse_pdf(path: Path) -> Tuple[str, str, None]:
    """
    PDF 三级引擎回退：PyMuPDF → pdfplumber → PyPDF2

    每级失败自动降级到下一级，全部失败才报错。
    图片扫描件（无可提取文本）在三轮尝试后会给出明确引导。
    """
    errors: list[str] = []

    # ── 第 1 级：PyMuPDF（最快、最准）──
    text = _try_pymupdf(path)
    if text:
        return text, "PDF", None
    errors.append("PyMuPDF 未能提取文本")

    # ── 第 2 级：pdfplumber（对复杂排版更好）──
    text = _try_pdfplumber(path)
    if text:
        return text, "PDF", None
    errors.append("pdfplumber 未能提取文本")

    # ── 第 3 级：PyPDF2（最兼容）──
    text = _try_pypdf2(path)
    if text:
        return text, "PDF", None
    errors.append("PyPDF2 未能提取文本")

    # 全部失败 → 可能是图片扫描件
    detail = " | ".join(errors)
    raise ParseError(
        f"PDF 解析失败（三级引擎均未提取到文本）: {detail}",
        "❌ 无法解析 PDF：三个解析引擎都未能提取到文字。"
        "PDF 可能是图片扫描件，请尝试直接粘贴简历文本。"
    )


def _try_pymupdf(path: Path) -> str:
    """PyMuPDF (fitz) 解析"""
    try:
        import fitz
    except ImportError:
        return ""

    try:
        doc = fitz.open(str(path))
        text_parts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                text_parts.append(text.strip())
        doc.close()
        return "\n\n".join(text_parts) if text_parts else ""
    except Exception:
        return ""


def _try_pdfplumber(path: Path) -> str:
    """pdfplumber 解析（对表格/复杂排版更好）"""
    try:
        import pdfplumber
    except ImportError:
        return ""

    try:
        with pdfplumber.open(str(path)) as pdf:
            text_parts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text.strip())
            return "\n\n".join(text_parts) if text_parts else ""
    except Exception:
        return ""


def _try_pypdf2(path: Path) -> str:
    """PyPDF2 解析（最广泛兼容）"""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return ""

    try:
        reader = PdfReader(str(path))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text.strip())
        return "\n\n".join(text_parts) if text_parts else ""
    except Exception:
        return ""


# ══════════════════════════════════════════════════════════════
# DOCX 解析
# ══════════════════════════════════════════════════════════════

def _parse_docx(path: Path) -> Tuple[str, str, str]:
    """解析 DOCX 文件，返回 (text, "DOCX", template_path)"""
    try:
        from docx import Document
    except ImportError:
        raise ParseError(
            "python-docx 未安装",
            "DOCX 解析组件未安装，请运行: pip install python-docx"
        )

    try:
        doc = Document(str(path))
        text_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # 也提取表格中的文字
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        if not text_parts:
            raise ParseError(
                "DOCX 无内容",
                "❌ DOCX 文件似乎没有内容，请检查文件是否正确。"
            )

        full_text = "\n".join(text_parts)
        return full_text, "DOCX", str(path)

    except ParseError:
        raise
    except Exception as e:
        raise ParseError(
            f"DOCX 解析异常: {str(e)}",
            f"❌ DOCX 文件解析失败，文件可能已损坏。请尝试直接粘贴简历文本。"
        )
